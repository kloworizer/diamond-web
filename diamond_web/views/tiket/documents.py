"""Tiket document generation — DOCX downloads for Tanda Terima, Lampiran, and Register Data."""

import re
from io import BytesIO

from django.contrib.auth.decorators import login_required, user_passes_test
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.utils import timezone
from django.views.decorators.http import require_GET

from ...models.detil_tanda_terima import DetilTandaTerima
from ...models.klasifikasi_jenis_data import KlasifikasiJenisData
from ...models.tiket import Tiket
from ...models.tiket_pic import TiketPIC


def _is_p3de_user(user):
    """Check if user is P3DE (can generate/download documents)."""
    if not user or not user.is_authenticated:
        return False
    if user.is_superuser or user.groups.filter(name='admin').exists():
        return True
    return user.groups.filter(name='user_p3de').exists()


def _format_periode_tiket(tiket_obj):
    """Format periode display for a tiket object using periode penerimaan."""
    if not tiket_obj.id_periode_data or not tiket_obj.id_periode_data.id_periode_pengiriman:
        return '-'

    periode_desc = tiket_obj.id_periode_data.id_periode_pengiriman.periode_penerimaan or '-'
    tahun = str(tiket_obj.tahun) if tiket_obj.tahun else '-'

    if str(periode_desc).lower() == 'bulanan' and tiket_obj.periode:
        bulan_map = {
            1: 'Januari', 2: 'Februari', 3: 'Maret', 4: 'April', 5: 'Mei', 6: 'Juni',
            7: 'Juli', 8: 'Agustus', 9: 'September', 10: 'Oktober', 11: 'November', 12: 'Desember'
        }
        bulan = bulan_map.get(tiket_obj.periode, f'Bulan {tiket_obj.periode}')
        return f"{bulan} {tahun}"
    if 'semester' in str(periode_desc).lower() and tiket_obj.periode:
        return f"Semester {tiket_obj.periode} {tahun}"
    if 'triwulan' in str(periode_desc).lower() and tiket_obj.periode:
        return f"Triwulan {tiket_obj.periode} {tahun}"
    if 'mingguan' in str(periode_desc).lower() and tiket_obj.periode:
        return f"Minggu {tiket_obj.periode} {tahun}"
    return f"{periode_desc} {tahun}"


def _safe_filename_part(raw):
    """Convert a string into a filename-safe format.

    Replaces all non-alphanumeric characters (except dot, underscore, hyphen) with
    underscores, strips leading/trailing underscores, and returns 'file' if the
    result is empty.

    Args:
        raw (str): The raw string to sanitize for use in filenames.

    Returns:
        str: A filename-safe string with special characters replaced by underscores.
             Returns 'file' if the input produces an empty string after sanitization.
    """
    return re.sub(r'[^A-Za-z0-9._-]+', '_', str(raw or '')).strip('_') or 'file'


def _format_date_indonesian(date_obj):
    """Format a date object as 'D bulan YYYY' in Indonesian.

    Converts a datetime.date object into Indonesian-formatted string with
    full month names in lowercase. For example: 4 januari 2026.

    Args:
        date_obj (datetime.date or None): Date object to format. If None or falsy,
                                         returns '-'.

    Returns:
        str: Formatted date string in Indonesian (e.g., '4 januari 2026') or '-'
             if date_obj is None/falsy.
    """
    if not date_obj:
        return '-'
    
    bulan_map = {
        1: 'januari', 2: 'februari', 3: 'maret', 4: 'april', 5: 'mei', 6: 'juni',
        7: 'juli', 8: 'agustus', 9: 'september', 10: 'oktober', 11: 'november', 12: 'desember'
    }
    bulan = bulan_map.get(date_obj.month, '')
    return f"{date_obj.day} {bulan} {date_obj.year}"


def _build_table_doc(title, headers, rows_data):
    """Build a DOCX document with a single table.

    Creates a Word document with a title heading and a formatted table. The table
    includes header row with column titles and data rows from the provided rows_data.
    Table style is set to 'Table Grid' for standard appearance.

    Args:
        title (str): Title to display as heading (level 1) in the document.
        headers (list of str): Column header labels for the table.
        rows_data (list of list): List of row data, where each row is a list of
                                 values that will be converted to strings.

    Returns:
        docx.Document: A Document object containing the title and formatted table,
                      ready to be saved to file or returned as HTTP response.
    """
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in rows_data:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
    return doc


@login_required
@user_passes_test(lambda u: _is_p3de_user(u))
@require_GET
def tiket_documents_download(request, pk):
    """Generate and download a single DOCX document for a tiket.

    Accepts a ``?doc_type=`` query parameter to select which document to return:

    - ``tanda_terima`` (default) — Tanda Terima Data
    - ``lampiran``               — Lampiran Tanda Terima
    - ``register``               — Register Data

    Access is restricted to admins/superusers and assigned TiketPIC members.
    Returns HTTP 400 if tanda_terima has not been created yet.
    """
    try:
        from docx import Document
    except ImportError:
        return HttpResponse('Library python-docx belum terpasang.', status=500)

    tiket = get_object_or_404(
        Tiket.objects.select_related(
            'id_periode_data__id_sub_jenis_data_ilap__id_ilap__id_kategori',
            'id_periode_data__id_sub_jenis_data_ilap__id_ilap__id_kpp__id_kanwil',
            'id_periode_data__id_sub_jenis_data_ilap__id_ilap',
            'id_periode_data__id_periode_pengiriman',
            'id_bentuk_data',
            'id_cara_penyampaian',
        ),
        pk=pk,
    )

    if not request.user.groups.filter(name='admin').exists() and not request.user.is_superuser:
        has_access = TiketPIC.objects.filter(id_tiket=tiket, id_user=request.user, active=True).exists()
        if not has_access:
            return HttpResponse('Tidak memiliki akses ke tiket ini.', status=403)

    if not tiket.tanda_terima:
        return HttpResponse('Dokumen hanya tersedia jika tanda terima sudah dibuat.', status=400)

    # ------------------------------------------------------------------ #
    # Collect tanda-terima group and associated tiket rows                #
    # ------------------------------------------------------------------ #
    detil = DetilTandaTerima.objects.select_related('id_tanda_terima').filter(id_tiket=tiket).order_by('-id').first()
    tanda_terima = detil.id_tanda_terima if detil else None

    if tanda_terima:
        tiket_ids = list(
            DetilTandaTerima.objects.filter(id_tanda_terima=tanda_terima).values_list('id_tiket_id', flat=True)
        )
        tiket_rows = list(
            Tiket.objects.filter(id__in=tiket_ids).select_related(
                'id_periode_data__id_sub_jenis_data_ilap__id_ilap',
                'id_periode_data__id_periode_pengiriman',
                'id_bentuk_data',
                'id_cara_penyampaian',
            ).order_by('id')
        )
    else:
        tiket_rows = [tiket]

    # ------------------------------------------------------------------ #
    # Dasar hukum lookup                                                  #
    # ------------------------------------------------------------------ #
    jenis_data_ids = {
        t.id_periode_data.id_sub_jenis_data_ilap_id
        for t in tiket_rows
        if t.id_periode_data and t.id_periode_data.id_sub_jenis_data_ilap_id
    }
    dasar_hukum_map = {}
    for row in KlasifikasiJenisData.objects.filter(id_jenis_data_ilap_id__in=jenis_data_ids).select_related('id_klasifikasi_tabel'):
        dasar_hukum_map.setdefault(row.id_jenis_data_ilap_id, []).append(row.id_klasifikasi_tabel.deskripsi)

    # ------------------------------------------------------------------ #
    # PIC P3DE name                                                       #
    # ------------------------------------------------------------------ #
    p3de = TiketPIC.objects.select_related('id_user').filter(
        id_tiket=tiket,
        role=TiketPIC.Role.P3DE,
        active=True,
    ).order_by('id').first()
    p3de_name = '-'
    if p3de and p3de.id_user:
        p3de_name = p3de.id_user.get_full_name().strip() or p3de.id_user.username

    # ------------------------------------------------------------------ #
    # Derived fields                                                      #
    # ------------------------------------------------------------------ #
    ilap = (
        tiket.id_periode_data.id_sub_jenis_data_ilap.id_ilap
        if tiket.id_periode_data and tiket.id_periode_data.id_sub_jenis_data_ilap
        else None
    )
    kategori_name = ((ilap.id_kategori.nama_kategori if ilap and ilap.id_kategori else '') or '').lower()
    if 'regional' in kategori_name and ilap and ilap.id_kpp and ilap.id_kpp.id_kanwil:
        diterima_dari = ilap.id_kpp.id_kanwil.nama_kanwil
    else:
        diterima_dari = ilap.nama_ilap if ilap else '-'

    # Collect multi-value fields from tiket_rows (deduplicated, like periode_data)
    periode_list = []
    nomor_surat_list = []
    tanggal_surat_list = []
    bentuk_data_list = []
    cara_penyampaian_list = []
    
    seen_periode = set()
    seen_nomor_surat = set()
    seen_tanggal_surat = set()
    seen_bentuk_data = set()
    seen_cara_penyampaian = set()
    
    for t in tiket_rows:
        # Periode
        label = _format_periode_tiket(t)
        if label not in seen_periode:
            seen_periode.add(label)
            periode_list.append(label)
        
        # Nomor Surat Pengantar
        nomor_surat = t.nomor_surat_pengantar or '-'
        if nomor_surat not in seen_nomor_surat:
            seen_nomor_surat.add(nomor_surat)
            nomor_surat_list.append(nomor_surat)
        
        # Tanggal Surat Pengantar
        tanggal = _format_date_indonesian(t.tanggal_surat_pengantar) if t.tanggal_surat_pengantar else '-'
        if tanggal not in seen_tanggal_surat:
            seen_tanggal_surat.add(tanggal)
            tanggal_surat_list.append(tanggal)
        
        # Bentuk Data
        bentuk = t.id_bentuk_data.deskripsi if t.id_bentuk_data else '-'
        if bentuk not in seen_bentuk_data:
            seen_bentuk_data.add(bentuk)
            bentuk_data_list.append(bentuk)
        
        # Cara Penyampaian
        cara = t.id_cara_penyampaian.deskripsi if t.id_cara_penyampaian else '-'
        if cara not in seen_cara_penyampaian:
            seen_cara_penyampaian.add(cara)
            cara_penyampaian_list.append(cara)

    nomor_tanda_terima = tanda_terima.nomor_tanda_terima_format if tanda_terima else '-'
    tgl_terima_dip = _format_date_indonesian(tiket.tgl_terima_dip) if tiket.tgl_terima_dip else '-'

    # ------------------------------------------------------------------ #
    # DOC 1 — Tanda Terima                                                #
    # ------------------------------------------------------------------ #
    doc_tanda = Document()
    doc_tanda.add_heading('Tanda Terima Data', level=1)
    fields = [
        ('Nomor Tanda Terima',    nomor_tanda_terima),
        ('Diterima Dari',         diterima_dari),
        ('Nomor Surat Pengantar',  ', '.join(nomor_surat_list) if nomor_surat_list else '-'),
        ('Tanggal Surat Pengantar', ', '.join(tanggal_surat_list) if tanggal_surat_list else '-'),
        ('Nama ILAP',             ilap.nama_ilap if ilap else '-'),
        ('Jenis Data',            'Terlampir'),
        ('Periode Data',          ', '.join(periode_list) if periode_list else '-'),
        ('Bentuk Data',           ', '.join(bentuk_data_list) if bentuk_data_list else '-'),
        ('Tanggal Terima DIP',    tgl_terima_dip),
        ('Cara Penyampaian',      ', '.join(cara_penyampaian_list) if cara_penyampaian_list else '-'),
        ('Nama PIC P3DE',         p3de_name),
    ]
    table_fields = doc_tanda.add_table(rows=0, cols=2)
    table_fields.style = 'Table Grid'
    for key, value in fields:
        row = table_fields.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    # ------------------------------------------------------------------ #
    # DOC 2 & 3 — Lampiran Tanda Terima / Register Data (shared rows)    #
    # ------------------------------------------------------------------ #
    lampiran_headers = ['Nama ILAP', 'Jenis Data', 'Periode Data Tahun', 'Baris Diterima', 'Dasar Hukum']
    lampiran_rows = []
    for t in tiket_rows:
        sub = t.id_periode_data.id_sub_jenis_data_ilap if t.id_periode_data else None
        ilap_obj = sub.id_ilap if sub else None
        dasar_hukum_list = dasar_hukum_map.get(sub.id, []) if sub else []
        lampiran_rows.append([
            f"{ilap_obj.id_ilap} - {ilap_obj.nama_ilap}" if ilap_obj else '-',
            f"{sub.id_sub_jenis_data} - {sub.nama_sub_jenis_data}" if sub else '-',
            _format_periode_tiket(t),
            str(t.baris_diterima if t.baris_diterima is not None else '-'),
            ', '.join(dasar_hukum_list) if dasar_hukum_list else '-',
        ])

    doc_lampiran = _build_table_doc('Lampiran Tanda Terima', lampiran_headers, lampiran_rows)
    doc_register = _build_table_doc('Register Data', lampiran_headers, lampiran_rows)

    # ------------------------------------------------------------------ #
    # Select document based on doc_type param and return response         #
    # ------------------------------------------------------------------ #
    now_ts = timezone.now().strftime('%Y%m%d_%H%M%S_%f')
    nomor_safe = _safe_filename_part(nomor_tanda_terima)
    doc_type = request.GET.get('doc_type', 'tanda_terima')

    if doc_type == 'lampiran':
        doc = doc_lampiran
        filename = f"lampiran_tanda_terima_{nomor_safe}_{now_ts}.docx"
    elif doc_type == 'register':
        doc = doc_register
        filename = f"register_data_{nomor_safe}_{now_ts}.docx"
    else:
        doc = doc_tanda
        filename = f"tanda_terima_{nomor_safe}_{now_ts}.docx"

    buffer = BytesIO()
    doc.save(buffer)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response
