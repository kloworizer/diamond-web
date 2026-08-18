"""Utility functions for DOCX template processing and placeholder replacement.

Two kinds of placeholders are supported:

  Simple placeholders  {{variable_name}}
      Replaced once with the matching value from the ``replacements`` dict.
      Keys in ``replacements`` must include the braces, e.g. ``'{{nama_ilap}}'``.

  Row placeholders     {{row.field_name}}
      Used in a table row to mark it as a *repeating template row*.
      The row is cloned once for every item in the ``row_data`` list and each
      ``{{row.field}}`` is replaced with ``item_dict['field']``.
      If ``row_data`` is empty the template row is silently removed.

Example template table::

    | No | Nama ILAP          | Jenis Data          | Periode          | Baris    |
    |----|--------------------|---------------------|------------------|----------|
    | {{row.no}} | {{row.nama_ilap}} | {{row.jenis_data}} | {{row.periode}} | {{row.baris}} |

Example row_data::

    row_data = [
        {'no': '1', 'nama_ilap': 'ILAP A', 'jenis_data': 'Bulanan', ...},
        {'no': '2', 'nama_ilap': 'ILAP B', 'jenis_data': 'Tahunan', ...},
    ]
"""

import logging
import re
from copy import deepcopy
from io import BytesIO
from docx import Document
from docx.table import _Row as DocxRow

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _iter_run_groups(paragraph):
    """Yield each contiguous group of Run objects in *paragraph*.

    A paragraph's own direct runs are one group; each hyperlink nested in the
    paragraph is a separate group. python-docx only exposes ``<w:r>`` elements
    that are direct children of ``<w:p>`` via ``paragraph.runs`` — runs living
    inside a ``<w:hyperlink>`` (a common leftover from pasted text) are
    invisible to that property, so without this a placeholder wrapped in a
    hyperlink is silently left unreplaced.
    """
    if paragraph.runs:
        yield paragraph.runs
    for hyperlink in paragraph.hyperlinks:
        if hyperlink.runs:
            yield hyperlink.runs


def _replace_in_runs(runs, new_text_fn):
    """Collapse a run group to a single run holding the replaced text.

    ``new_text_fn`` receives the concatenated text of *runs* and returns the
    replacement text, or ``None`` if nothing in this group needs replacing.
    Rewriting ``runs[0].text`` keeps that run's own formatting (``<w:rPr>``)
    untouched; the remaining runs in the group are dropped.
    """
    full_text = ''.join(run.text for run in runs)
    new_text = new_text_fn(full_text)
    if new_text is None:
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run._element.getparent().remove(run._element)


def _replace_in_paragraph(paragraph, replacements):
    """Replace all simple {{key}} placeholders in a paragraph while preserving formatting.

    Scans the paragraph's direct runs and any hyperlink-nested runs separately.
    """
    def resolve(full_text):
        if not any(ph in full_text for ph in replacements):
            return None
        new_text = full_text
        for placeholder, value in replacements.items():
            new_text = new_text.replace(placeholder, str(value if value is not None else '-'))
        return new_text

    for runs in _iter_run_groups(paragraph):
        _replace_in_runs(runs, resolve)


_ROW_PLACEHOLDER_RE = re.compile(r'\{\{row\.\s*\w+')


def _row_has_row_placeholder(row):
    """Return True if any cell in *row* contains a ``{{row.xxx}}`` placeholder."""
    for cell in row.cells:
        for para in cell.paragraphs:
            # paragraph.text includes hyperlink-nested run text, unlike paragraph.runs.
            if _ROW_PLACEHOLDER_RE.search(para.text):
                return True
    return False


def _fill_row_placeholders(row, row_dict):
    """Replace ``{{row.field}}`` placeholders in all paragraphs of *row* while preserving formatting."""
    def resolve(full_text):
        if not _ROW_PLACEHOLDER_RE.search(full_text):
            return None
        new_text = full_text
        for key, value in row_dict.items():
            pattern = r'\{\{row\.\s*' + re.escape(key) + r'\s*\}\}'
            new_text = re.sub(pattern, str(value if value is not None else '-'), new_text)
        return new_text

    for cell in row.cells:
        for para in cell.paragraphs:
            for runs in _iter_run_groups(para):
                _replace_in_runs(runs, resolve)


def _expand_repeating_rows(table, row_data):
    """Expand every template row (one containing ``{{row.xxx}}``) in *table*.

    For each template row found:
    - If ``row_data`` is non-empty, clone the row once per item and fill
      ``{{row.field}}`` placeholders from the item dict.
    - The original template row is always removed afterwards.
    """
    # Collect indices of template rows (scan all rows once)
    template_indices = [
        i for i, row in enumerate(table.rows)
        if _row_has_row_placeholder(row)
    ]
    if not template_indices:
        return

    # Process from the last index to the first so earlier indices stay valid
    for idx in reversed(template_indices):
        template_tr = table.rows[idx]._element
        parent = template_tr.getparent()

        if row_data:
            # Insert cloned rows right after the template row, in order
            insert_after = template_tr
            for item_dict in row_data:
                new_tr = deepcopy(template_tr)
                insert_after.addnext(new_tr)
                insert_after = new_tr
                _fill_row_placeholders(DocxRow(new_tr, table), item_dict)

        # Remove the original template row
        parent.remove(template_tr)


def _iter_nested_tables(table):
    """Yield *table* and all nested tables inside its cells."""
    yield table
    for row in table.rows:
        for cell in row.cells:
            for nested in cell.tables:
                yield from _iter_nested_tables(nested)


def _iter_all_tables(doc):
    """Yield all tables from document body, headers, and footers (including nested tables)."""
    seen = set()

    def _yield_unique(table):
        table_id = id(table._element)
        if table_id in seen:
            return
        seen.add(table_id)
        yield table

    # Body tables
    for table in doc.tables:
        for nested_table in _iter_nested_tables(table):
            yield from _yield_unique(nested_table)

    # Header/footer tables per section
    for section in doc.sections:
        for table in section.header.tables:
            for nested_table in _iter_nested_tables(table):
                yield from _yield_unique(nested_table)
        for table in section.footer.tables:
            for nested_table in _iter_nested_tables(table):
                yield from _yield_unique(nested_table)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def fill_template_with_data(template_file, replacements, row_data=None):
    """Fill a DOCX template with data by replacing placeholders.

    Args:
        template_file: File path string or file-like object of the template DOCX.
        replacements:  Dict of ``{'{{key}}': value}`` for simple placeholders.
        row_data:      Optional list of dicts for repeating table rows.
                       Each dict maps field names to values used in
                       ``{{row.field}}`` placeholders.

    Returns:
        BytesIO: The filled document ready for download.
    """
    doc = Document(template_file)

    # 1. Body paragraphs — simple replacements
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # 2. Tables — first expand repeating rows, then simple replacements
    # Includes body tables, header/footer tables, and nested tables.
    for table in _iter_all_tables(doc):
        _expand_repeating_rows(table, row_data)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    # 3. Headers and footers — simple replacements only
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, replacements)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
